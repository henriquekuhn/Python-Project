from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_section(doc, title, content):
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

def add_subsection(doc, title, content):
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

def add_table(doc, data, column_names):
    table = doc.add_table(rows=1, cols=len(column_names))
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(column_names):
        hdr_cells[i].text = column_name
    for row in data:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = cell_value

def add_table_of_contents(doc):
    doc.add_paragraph('Índice', style='Heading 1')
    toc = doc.add_paragraph('1. Alternativas de Banco de Dados e Vantagens do SQL', style='List Number')
    toc.add_run('\n    1.1 Bancos de Dados Relacionais').bold = True
    toc.add_run('\n    1.2 Bancos de Dados Não Relacionais').bold = True
    toc.add_run('\n    1.3 Vantagens dos Bancos de Dados Relacionais').bold = True
    toc.add_run('\n2. Escolha do SQLite com SQLAlchemy').bold = True
    toc.add_run('\n    2.1 Descrição do SQLite').bold = True
    toc.add_run('\n    2.2 Vantagens do SQLite com SQLAlchemy').bold = True
    toc.add_run('\n    2.3 Limitações do SQLite').bold = True
    toc.add_run('\n    2.4 Explicação do ORM').bold = True
    toc.add_run('\n3. Escolha do MariaDB').bold = True
    toc.add_run('\n    3.1 Descrição do MariaDB').bold = True
    toc.add_run('\n    3.2 Vantagens do MariaDB').bold = True
    toc.add_run('\n    3.3 Limitações do MariaDB').bold = True
    toc.add_run('\n4. Explicação dos Métodos Principais').bold = True
    toc.add_run('\n    4.1 Métodos do SQLite').bold = True
    toc.add_run('\n    4.2 Métodos do MariaDB').bold = True
    toc.add_run('\n5. Comparação Final').bold = True
    toc.add_run('\n6. Conclusão').bold = True

def add_code_block(doc, code, language='python'):
    code_paragraph = doc.add_paragraph()
    code_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    code_run = code_paragraph.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10.5)
    code_run.font.color.rgb = RGBColor(255, 255, 255)  # Texto branco

    # Formatação para cores de sintaxe básica
    for run in code_paragraph.runs:
        if 'def ' in run.text or 'class ' in run.text:
            run.font.color.rgb = RGBColor(86, 156, 214)  # Azul claro
        elif '(' in run.text or ')' in run.text or ',' in run.text:
            run.font.color.rgb = RGBColor(220, 220, 170)  # Amarelo claro
        elif '"' in run.text or "'" in run.text:
            run.font.color.rgb = RGBColor(214, 157, 133)  # Marrom claro

doc = Document()

doc.add_heading('Documentação Técnica: Comparação entre SQLite e MariaDB', 0)

add_table_of_contents(doc)

add_section(doc, '1. Alternativas de Banco de Dados e Vantagens do SQL',
    'Ao considerar a escolha de um banco de dados para o desenvolvimento de aplicações, é essencial avaliar diversas alternativas '
    'e suas características. Entre as opções mais comuns estão bancos de dados relacionais e não relacionais. Neste documento, '
    'focaremos nas vantagens dos bancos de dados relacionais, como o SQLite e o MariaDB, e exploraremos outras alternativas.')

add_subsection(doc, '1.1 Bancos de Dados Relacionais',
    'Os bancos de dados relacionais utilizam um modelo de dados baseado em tabelas que estão inter-relacionadas. A Structured Query Language (SQL) é usada para manipular e consultar os dados. '
    'Este modelo proporciona integridade referencial, suporte para transações complexas e um alto grau de normalização dos dados. '
    'Além disso, bancos de dados relacionais suportam a criação de índices para otimizar a velocidade das consultas.')

add_subsection(doc, '1.2 Bancos de Dados Não Relacionais',
    'Os bancos de dados não relacionais, como MongoDB e Redis, são projetados para atender a necessidades específicas que não se encaixam bem no modelo relacional. '
    'Eles oferecem flexibilidade na modelagem dos dados e são escaláveis horizontalmente, o que pode ser vantajoso em aplicações que requerem alta disponibilidade e desempenho.')

add_subsection(doc, '1.3 Vantagens dos Bancos de Dados Relacionais',
    '1. **Integridade dos Dados:** Bancos de dados relacionais garantem que os dados sejam armazenados e manipulados de forma consistente e confiável.\n'
    '2. **Consultas Poderosas:** A SQL permite a realização de consultas complexas e análises detalhadas.\n'
    '3. **Transações ACID:** Garantia de propriedades de Atomicidade, Consistência, Isolamento e Durabilidade.\n'
    '4. **Normalização:** Reduz a redundância dos dados e melhora a integridade dos mesmos.')

add_table(doc, [
    ['Tipo de Banco', 'Descrição', 'Vantagens', 'Desvantagens'],
    ['SQLite', 'Banco de dados relacional leve e embutido.', 'Facilidade de configuração e uso, baixo overhead.', 'Limitações em escalabilidade e concorrência.'],
    ['MariaDB', 'Sistema de gerenciamento de banco de dados relacional completo.', 'Alta escalabilidade, suporte a replicação e particionamento.', 'Requer configuração e manutenção.'],
    ['PostgreSQL', 'Banco de dados relacional avançado e open-source.', 'Recursos avançados, suporte a grandes volumes de dados.', 'Complexidade na configuração.'],
    ['MongoDB', 'Banco de dados NoSQL orientado a documentos.', 'Flexibilidade na modelagem de dados, escalabilidade.', 'Não relacional, menos suporte a transações complexas.'],
    ['Redis', 'Banco de dados NoSQL em memória.', 'Alta performance em leitura, ideal para cache.', 'Persistência limitada, não ideal para dados críticos.']
], ['Banco de Dados', 'Descrição', 'Vantagens', 'Desvantagens'])

add_section(doc, '2. Escolha do SQLite com SQLAlchemy',
    'O SQLite foi escolhido para o projeto inicial devido à sua simplicidade e facilidade de uso. O SQLAlchemy foi utilizado para fornecer uma camada de abstração sobre o banco de dados.')

add_subsection(doc, '2.1 Descrição do SQLite',
    'O SQLite é um banco de dados relacional embutido que armazena os dados em um único arquivo. É conhecido por sua leveza e facilidade de integração em aplicações. '
    'Não requer um servidor separado, o que simplifica a configuração e reduz os requisitos de manutenção.')

add_subsection(doc, '2.2 Vantagens do SQLite com SQLAlchemy',
    '1. **Facilidade de Uso:** A configuração e o gerenciamento do SQLite são simples, tornando-o ideal para desenvolvimento e protótipos.\n'
    '2. **Leveza:** O SQLite é extremamente leve, com baixo overhead de recursos.\n'
    '3. **Portabilidade:** Os dados são armazenados em um único arquivo, facilitando a transferência entre sistemas.\n'
    '4. **Integração com SQLAlchemy:** O SQLAlchemy fornece uma camada ORM que simplifica a interação com o banco de dados.')

add_subsection(doc, '2.3 Limitações do SQLite',
    '1. **Escalabilidade:** Não é ideal para aplicações que requerem alta escalabilidade e suporte a muitos usuários simultâneos.\n'
    '2. **Funcionalidades Avançadas:** Falta suporte para algumas funcionalidades avançadas encontradas em outros sistemas de banco de dados relacionais.\n'
    '3. **Segurança:** Em ambientes multiusuário, pode não oferecer o mesmo nível de segurança e controle de acesso que outros bancos de dados.')

add_subsection(doc, '2.4 Explicação do ORM',
    'O Object-Relational Mapping (ORM) é uma técnica que permite interagir com um banco de dados relacional usando uma linguagem de programação orientada a objetos. '
    'O SQLAlchemy é uma biblioteca de ORM para Python que facilita a manipulação de bancos de dados relacionais de forma programática. Com o ORM, você pode definir modelos de dados como classes Python, '
    'e o SQLAlchemy traduz automaticamente essas classes em tabelas e consultas SQL.')

add_section(doc, '3. Escolha do MariaDB',
    'MariaDB foi escolhido como a solução para a produção devido à sua robustez e escalabilidade. Oferece uma série de recursos avançados que atendem a necessidades de alta demanda.')

add_subsection(doc, '3.1 Descrição do MariaDB',
    'MariaDB é um sistema de gerenciamento de banco de dados relacional que surgiu como uma bifurcação do MySQL. É conhecido por sua alta performance, escalabilidade e suporte a grandes volumes de dados. '
    'Oferece funcionalidades avançadas como replicação, particionamento de tabelas e suporte a transações complexas.')

add_subsection(doc, '3.2 Vantagens do MariaDB',
    '1. **Alta Escalabilidade:** Suporte para grandes volumes de dados e alta concorrência.\n'
    '2. **Funcionalidades Avançadas:** Inclui replicação, particionamento de tabelas e suporte a transações complexas.\n'
    '3. **Desempenho:** Desempenho superior em cargas de trabalho intensivas.')

add_subsection(doc, '3.3 Limitações do MariaDB',
    '1. **Complexidade de Configuração:** Requer mais configuração e manutenção em comparação com soluções mais simples como o SQLite.\n'
    '2. **Overhead de Recursos:** Maior consumo de recursos em comparação com bancos de dados embutidos.')

add_section(doc, '4. Explicação dos Métodos Principais',
    'Os métodos principais para interagir com bancos de dados SQLite e MariaDB são apresentados a seguir, com uma breve explicação de suas funcionalidades.')

add_subsection(doc, '4.1 Métodos do SQLite',
    'Para o SQLite, os principais métodos são:\n')

# Adicionando o código formatado do SQLite
sqlite_code = '''def connect_to_db(db_name):
    """Estabelece uma conexão com o banco de dados SQLite."""
    # Código para conectar ao banco de dados SQLite
    pass

def create_table(conn, table_name, columns):
    """Cria uma tabela no banco de dados SQLite."""
    # Código para criar uma tabela
    pass

def insert_data(conn, name, email):
    """Insere dados na tabela de usuários."""
    # Código para inserir dados
    pass
'''

add_code_block(doc, sqlite_code)

add_subsection(doc, '4.2 Métodos do MariaDB',
    'Para o MariaDB, os principais métodos são:\n')

# Adicionando o código formatado do MariaDB
mariadb_code = '''def connect_engine(user, password, host, port):
    """Estabelece uma conexão com o banco de dados MariaDB usando as credenciais fornecidas."""
    # Código para conectar ao banco de dados MariaDB
    pass

def check_database(engine, db_name):
    """Verifica se o banco de dados especificado existe."""
    # Código para verificar banco de dados
    pass

def create_database(engine, db_name):
    """Cria um banco de dados com o nome fornecido."""
    # Código para criar banco de dados
    pass

def create_table(engine, db_name, table_name, columns):
    """Cria uma tabela no banco de dados especificado."""
    # Código para criar tabela
    pass

def insert_data(engine, db_name, table_name, data):
    """Insere dados na tabela do banco de dados especificado."""
    # Código para inserir dados
    pass
'''

add_code_block(doc, mariadb_code)

add_section(doc, '5. Comparação Final',
    'A escolha entre SQLite e MariaDB deve ser baseada nas necessidades específicas do projeto:\n'
    '1. **Escalabilidade:** O MariaDB é mais adequado para aplicações que exigem escalabilidade e suporte a muitos usuários simultâneos.\n'
    '2. **Facilidade de Uso:** O SQLite é ideal para desenvolvimento rápido e protótipos, devido à sua simplicidade e baixo overhead.\n'
    '3. **Funcionalidades Avançadas:** O MariaDB oferece mais funcionalidades avançadas, como replicação e particionamento de tabelas.\n'
    '4. **Performance:** O SQLite é eficiente para cargas de trabalho leves, enquanto o MariaDB é melhor para cargas intensivas.\n'
    '5. **Recursos:** O MariaDB requer mais recursos e configuração, mas oferece melhor suporte para ambientes de produção complexos.')

add_section(doc, '6. Conclusão',
    'Ambos os bancos de dados, SQLite e MariaDB, têm suas próprias vantagens e desvantagens. A escolha depende das necessidades específicas do projeto, como escalabilidade, facilidade de uso, '
    'funcionalidades avançadas e performance. Para desenvolvimento rápido e prototipagem, o SQLite é uma excelente escolha devido à sua simplicidade. No entanto, para aplicações em produção que exigem '
    'alta escalabilidade e suporte a funcionalidades avançadas, o MariaDB é mais adequado.')

# Salva o documento
doc.save('C:/Users/adm_cafrunikuhn/Desktop/Database_Comparison_Document.docx')
