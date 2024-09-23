from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_word_document():
    doc = Document()

    # Define document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(0.5)
        section.bottom_margin = Pt(0.5)
        section.left_margin = Pt(0.5)
        section.right_margin = Pt(0.5)

    # Define custom style for code
    styles = doc.styles
    code_style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10.5)
    code_style.font.color.rgb = RGBColor(255, 255, 255)
    code_style.paragraph_format.alignment = 1  # Center alignment

    # Add background color to paragraphs using custom style
    def set_background_color(paragraph, color):
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        pPr.append(shd)

    def add_heading(doc, text, level=1):
        doc.add_heading(text, level=level)

    def add_paragraph(doc, text, style=None):
        doc.add_paragraph(text, style=style)

    def add_section(doc, title, subtitle, code, explanation):
        add_heading(doc, title, level=2)
        add_paragraph(doc, subtitle, style='Heading2')
        p = doc.add_paragraph(code, style='CodeStyle')
        set_background_color(p, '000000')  # Black background
        add_paragraph(doc, explanation)

    # SQLite code and explanation
    sqlite_code = '''
import sqlalchemy
from sqlalchemy.orm import declarative_base, sessionmaker
from sqlalchemy import Integer, Float, String, Column, desc
import os

#DIR = os.path.join(os.getcwd(), "My Database SQLAlchemy")
DIR = "C:\\Users\\adm_cafrunikuhn\\Desktop\\Henrique\\Repositories\\Python-Codes\\My Database SQLAlchemy"
print(DIR)
print(os.getcwd())
if os.getcwd() != DIR:
    os.chdir(DIR)

# Connecting to the Database
FILE_NAME = 'tests_database.db'
DATABASE_FILENAME = 'sqlite:///tests_database.db'
engine = sqlalchemy.create_engine(DATABASE_FILENAME, echo=True)

# Mapping Declaration
Base = declarative_base()
class Tests(Base):
    __tablename__ = 'tests'  # Required
    id = Column(Integer, primary_key=True)  # Required
    test_name = Column(String(50))
    test_temp = Column(Float)
    test_voltage = Column(Float)
    low_condition = Column(Float)
    high_condition = Column(Float)
    test_result = Column(Float)

class DatabaseManagement:
    def create_db(self):
        # Create the table in the my_database database
        ## Check if the database exists
        self.db_exists = os.path.exists(os.path.join(DIR, FILE_NAME))
        if not self.db_exists:
            print("Creating my_database database...")
            Base.metadata.create_all(engine)

    def initiate_db_session(self):
        # Create database session
        Session = sessionmaker(bind=engine)
        self.session = Session()

    def create_test(self, test_name, test_temp, test_voltage, low_condition, high_condition, test_result):    
        tests = []
        tests.append(Tests(test_name=test_name, test_temp=test_temp, test_voltage=test_voltage,
            low_condition=low_condition, high_condition=high_condition, test_result=test_result))

        self.session.add_all(tests)
        self.session.commit()

        return "Test create success."

    def query_test(self):
        # Query the database
        tests_db = self.session.query(Tests).all()
        db_size = len(tests_db)
        return db_size

    def query_by_filter(self, column_search, query_filter):  
        query_filter = query_filter
        column_search = column_search  
        column = getattr(Tests, column_search)
        tests_db = self.session.query(Tests).filter(column == query_filter).order_by(desc(Tests.id)).first()

        return getattr(tests_db, column_search)

    def edit_register(self, column_search, query_item, new_value):
        column = getattr(Tests, column_search)
        edit_db = self.session.query(Tests).filter(column == query_item).order_by(desc(Tests.id)).first()       
        edit_db.test_name = new_value
        self.session.commit()
        new_value_db = self.session.query(Tests).filter(column == new_value).order_by(desc(Tests.id)).first()
        return getattr(new_value_db, column_search)
    
    def delete_regiter(self, column_search, query_item):
        column = getattr(Tests, column_search)
        delete_db_list = self.session.query(Tests).filter(column == query_item).all()
        if delete_db_list:
            for delete_db in delete_db_list:
                self.session.delete(delete_db)

            self.session.commit()        
            deleted_db = self.session.query(Tests).filter(column == query_item).order_by(desc(Tests.id)).first()
            return deleted_db
        else:
            return "No items found matching the provided criteria."

        return deleted_db
'''

    sqlite_code_explanation = '''
    Este código é um exemplo de como configurar e usar um banco de dados SQLite com SQLAlchemy. 

    - `import sqlalchemy`: Importa o módulo SQLAlchemy para gerenciamento de bancos de dados.
    - `from sqlalchemy.orm import declarative_base, sessionmaker`: Importa classes para criar modelos e gerenciar sessões.
    - `from sqlalchemy import Integer, Float, String, Column, desc`: Importa tipos de dados e funções para definir colunas e ordenação.
    - `import os`: Importa o módulo os para manipulação de caminhos e arquivos.
    
    - `DIR` e `FILE_NAME`: Define o diretório e o nome do arquivo do banco de dados.
    - `DATABASE_FILENAME`: URL de conexão com o banco de dados SQLite.
    - `engine = sqlalchemy.create_engine(DATABASE_FILENAME, echo=True)`: Cria um mecanismo de conexão com o banco de dados.

    - `Base = declarative_base()`: Cria uma classe base para os modelos.
    - `class Tests(Base)`: Define a tabela `tests` no banco de dados com colunas específicas.
    - `class DatabaseManagement`: Classe para gerenciar operações no banco de dados, como criação, consulta, edição e exclusão de registros.
    - Métodos incluem `create_db`, `initiate_db_session`, `create_test`, `query_test`, `query_by_filter`, `edit_register`, e `delete_regiter`.
    '''

    mariadb_code = '''
import pymysql
import os
from dotenv import load_dotenv
from contextlib import closing, contextmanager

# Load environment variables
load_dotenv()

# Retrieve database credentials from environment variables
db_user = os.getenv('DB_USER')
db_password = os.getenv('DB_PASSWORD')
db_host = os.getenv('DB_HOST')

class DatabaseManagement:
    """
    A class to manage MySQL/MariaDB databases and tables.
    """

    def __init__(self):
        """
        Initialize DatabaseManagement class with None attributes for engine and connection.
        """
        self.connection = None

    @contextmanager
    def manage_connection(self):
        """
        Context manager to handle database connections.
        Yields:
            pymysql.connections.Connection: A pymysql connection object.
        """
        conn = None
        try:
            conn = pymysql.connect(host=db_host, user=db_user, password=db_password)
            yield conn
        finally:
            if conn:
                conn.close()

    def connect_engine(self):
        """
        Establishes a connection to the database using pymysql.

        Args:
            timeout_seconds (int): Timeout duration in seconds for the connection attempt.

        Returns:
            str: "Pass" if connection is successful, "Fail" if unsuccessful.
        """
        print("\nEstablishing connection: ")
        try:
            self.connection = pymysql.connect(host=db_host, user=db_user, password=db_password)
            print("Connection established successfully.")
            return "Pass"
        except pymysql.MySQLError as e:
            print(f"Error trying to connect to the database: {e}")
            return "Fail"

    def check_database(self, db_name):
        """
        Checks if a database exists with the given name.

        Args:
            db_name (str): The name of the database to check.

        Returns:
            tuple or None: A tuple representing the database if found, None if not found.
        """
        with self.manage_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SHOW DATABASES LIKE %s", (db_name,))
            result = cursor.fetchone()
            return result

    def create_database(self, db_name):
        """
        Creates a new database with the given name.

        Args:
            db_name (str): The name of the database to create.

        Returns:
            str: "Database created" if successful, "Fail" if unsuccessful.
        """
        with self.manage_connection() as conn:
            cursor = conn.cursor()
            try:
                cursor.execute(f"CREATE DATABASE {db_name}")
                return "Database created"
            except pymysql.MySQLError as e:
                print(f"Error creating database: {e}")
                return "Fail"

    def create_table(self, db_name, table_name, columns):
        """
        Creates a new table in the specified database with the given columns.

        Args:
            db_name (str): The name of the database.
            table_name (str): The name of the table to create.
            columns (dict): A dictionary of column names and their types.

        Returns:
            str: "Table created" if successful, "Fail" if unsuccessful.
        """
        with self.manage_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(f"USE {db_name}")
            columns_str = ", ".join([f"{col} {typ}" for col, typ in columns.items()])
            try:
                cursor.execute(f"CREATE TABLE {table_name} ({columns_str})")
                return "Table created"
            except pymysql.MySQLError as e:
                print(f"Error creating table: {e}")
                return "Fail"

    def insert_data(self, db_name, table_name, data):
        """
        Inserts data into the specified table in the given database.

        Args:
            db_name (str): The name of the database.
            table_name (str): The name of the table to insert data into.
            data (dict): A dictionary of column names and their values.

        Returns:
            str: "Data inserted" if successful, "Fail" if unsuccessful.
        """
        with self.manage_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(f"USE {db_name}")
            columns = ", ".join(data.keys())
            placeholders = ", ".join(['%s'] * len(data))
            try:
                cursor.execute(f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})", list(data.values()))
                return "Data inserted"
            except pymysql.MySQLError as e:
                print(f"Error inserting data: {e}")
                return "Fail"
'''

    mariadb_code_explanation = '''
    Este código é um exemplo de como configurar e usar um banco de dados MySQL/MariaDB com pymysql.

    - `import pymysql`: Importa o módulo pymysql para conexão com MySQL/MariaDB.
    - `import os`: Importa o módulo os para manipulação de variáveis de ambiente.
    - `from dotenv import load_dotenv`: Importa a função para carregar variáveis de ambiente de um arquivo .env.
    - `from contextlib import closing, contextmanager`: Importa o módulo para gerenciar recursos de forma eficiente.

    - `load_dotenv()`: Carrega as variáveis de ambiente a partir do arquivo .env.
    - `db_user`, `db_password`, `db_host`: Obtém as credenciais do banco de dados das variáveis de ambiente.
    
    - `class DatabaseManagement`: Define uma classe para gerenciar operações de banco de dados, como conexão, criação de banco de dados e tabelas, e inserção de dados.
    - Métodos incluem `connect_engine`, `check_database`, `create_database`, `create_table`, e `insert_data`.
    '''

    # Adding sections to the document
    add_section(doc, '1. SQLite', 'Código Python para SQLite', sqlite_code, sqlite_code_explanation)
    add_section(doc, '2. MariaDB', 'Código Python para MariaDB', mariadb_code, mariadb_code_explanation)

    # Save the document
    doc.save('Database_Comparison_Document.docx')

# Execute the function to create the document
create_word_document()
